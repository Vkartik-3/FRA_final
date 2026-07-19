#!/usr/bin/env python3
"""
Convert markdown semester report to Word document (.docx).
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import re

def parse_markdown_to_docx(md_file, docx_file):
    """
    Convert markdown to Word document with proper formatting.

    Args:
        md_file: Path to markdown file
        docx_file: Path to output Word document
    """
    # Create document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Read markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Main title (# )
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
            i += 1
            continue

        # Section headers (## )
        if line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_heading(text, level=2)
            run = p.runs[0]
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)  # Navy
            i += 1
            continue

        # Subsection headers (### )
        if line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_heading(text, level=3)
            run = p.runs[0]
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 102, 153)  # Blue
            i += 1
            continue

        # Subsubsection headers (#### )
        if line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_heading(text, level=4)
            run = p.runs[0]
            run.font.size = Pt(12)
            run.font.bold = True
            i += 1
            continue

        # Horizontal rules (---)
        if line.strip() == '---':
            doc.add_paragraph()  # Add space
            i += 1
            continue

        # Code blocks (```)
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip())
                i += 1

            # Add code block
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph(code_text)
            p.style = 'Normal'
            for run in p.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0, 0, 0)

            # Add shading
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F0F0F0')  # Light gray
            p._element.get_or_add_pPr().append(shading_elm)

            i += 1
            continue

        # Tables (| ... |)
        if line.strip().startswith('|') and '|' in line:
            # Collect table lines
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1

            # Parse table
            if len(table_lines) >= 2:
                # Extract header
                header_cells = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]

                # Skip separator line
                data_lines = table_lines[2:] if len(table_lines) > 2 else []

                # Create table
                table = doc.add_table(rows=1, cols=len(header_cells))
                table.style = 'Light Grid Accent 1'

                # Add header
                hdr_cells = table.rows[0].cells
                for idx, header in enumerate(header_cells):
                    hdr_cells[idx].text = header
                    for paragraph in hdr_cells[idx].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(10)

                # Add data rows
                for line in data_lines:
                    cells = [cell.strip() for cell in line.split('|')[1:-1]]
                    row_cells = table.add_row().cells
                    for idx, cell in enumerate(cells):
                        if idx < len(row_cells):
                            row_cells[idx].text = cell
                            for paragraph in row_cells[idx].paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)

                doc.add_paragraph()  # Add space after table

            continue

        # Bullet lists (- or *)
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
            i += 1
            continue

        # Numbered lists (1. )
        if re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line).strip()
            p = doc.add_paragraph(text, style='List Number')
            p.paragraph_format.left_indent = Inches(0.25)
            i += 1
            continue

        # Block quotes (> )
        if line.startswith('> '):
            text = line[2:].strip()
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Inches(0.5)
            for run in p.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(64, 64, 64)
            i += 1
            continue

        # Regular paragraphs
        text = line.strip()

        # Handle bold (**text**)
        text = re.sub(r'\*\*([^*]+)\*\*', r'<<<BOLD>>>\1<<<ENDBOLD>>>', text)

        # Handle italic (*text*)
        text = re.sub(r'\*([^*]+)\*', r'<<<ITALIC>>>\1<<<ENDITALIC>>>', text)

        # Handle inline code (`code`)
        text = re.sub(r'`([^`]+)`', r'<<<CODE>>>\1<<<ENDCODE>>>', text)

        p = doc.add_paragraph()

        # Parse formatting markers
        parts = re.split(r'(<<<BOLD>>>|<<<ENDBOLD>>>|<<<ITALIC>>>|<<<ENDITALIC>>>|<<<CODE>>>|<<<ENDCODE>>>)', text)

        bold_active = False
        italic_active = False
        code_active = False

        for part in parts:
            if part == '<<<BOLD>>>':
                bold_active = True
            elif part == '<<<ENDBOLD>>>':
                bold_active = False
            elif part == '<<<ITALIC>>>':
                italic_active = True
            elif part == '<<<ENDITALIC>>>':
                italic_active = False
            elif part == '<<<CODE>>>':
                code_active = True
            elif part == '<<<ENDCODE>>>':
                code_active = False
            elif part:
                run = p.add_run(part)
                run.font.bold = bold_active
                run.font.italic = italic_active
                if code_active:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(139, 0, 0)  # Dark red

        i += 1

    # Save document
    doc.save(docx_file)
    print(f"✅ Word document saved: {docx_file}")
    print(f"   File size: {Path(docx_file).stat().st_size / 1024:.1f} KB")


def main():
    script_dir = Path(__file__).parent
    md_file = script_dir / "FRA_Semester_Report_Fall_2025_FINAL.md"
    docx_file = script_dir / "FRA_Semester_Report_Fall_2025.docx"

    if not md_file.exists():
        print(f"❌ Error: Markdown file not found: {md_file}")
        return

    print(f"📄 Converting {md_file.name} to Word document...")
    parse_markdown_to_docx(md_file, docx_file)
    print(f"\n✅ Conversion complete!")
    print(f"   Input:  {md_file}")
    print(f"   Output: {docx_file}")


if __name__ == "__main__":
    main()
